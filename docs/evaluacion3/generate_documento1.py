#!/usr/bin/env python3
"""Genera Documento 1 (Criterios_Aceptacion_Epic4_Epic5.docx) para la Evaluacion 3.

Documento de ESPECIFICACION: define los 6 criterios de aceptacion en Gherkin
con su descripcion detallada, antes/independiente de la ejecucion de pruebas
(eso es el Documento 2). Reutiliza screenshots reales de la app (capturados
por la suite Playwright) para ilustrar cada escenario.
"""
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TESTS_DIR = Path(__file__).parent.parent.parent / "tests"
SCREENSHOTS = TESTS_DIR / "test-results"

PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
GRAY = RGBColor(0x52, 0x52, 0x5B)

CRITERIA = [
    {
        "epic": "Épica 4 — Reservas",
        "num": 1,
        "title": "Crear reserva exitosa con validaciones",
        "given": "Usuario autenticado en la plataforma",
        "when": "Selecciona un paquete disponible y completa la reserva con datos válidos",
        "then": "El sistema crea la reserva, decrementa cupos disponibles y muestra confirmación",
        "description": (
            "Cubre el flujo principal de la Épica 4: un cliente logueado elige un paquete con "
            "cupos disponibles, indica la cantidad de pasajeros dentro del rango permitido, y "
            "confirma. El sistema debe persistir la reserva en estado PENDING, descontar los "
            "cupos usados del paquete (BookingService, Regla 5), y llevar al usuario al paso de "
            "pago. Es el criterio base: si este falla, ningún otro flujo de compra es alcanzable."
        ),
        "screenshot": SCREENSHOTS / "epic4-reservas-Epica-4---R-16406-va-exitosa-con-validaciones-chromium" / "test-finished-1.png",
    },
    {
        "epic": "Épica 4 — Reservas",
        "num": 2,
        "title": "Aplicar descuentos automáticamente en reserva",
        "given": "Usuario es cliente frecuente (3+ reservas confirmadas)",
        "when": "Realiza una nueva reserva",
        "then": "El sistema aplica automáticamente descuento por cliente frecuente (10%)",
        "description": (
            "Verifica la Regla 2 de descuentos (DiscountService): el sistema cuenta las reservas "
            "CONFIRMED previas del cliente y, si son 3 o más, aplica un 10% de descuento sin que "
            "el cliente lo solicite ni pueda desactivarlo. El descuento debe quedar visible en el "
            "resumen de pago, con su porcentaje y descripción, antes de que el cliente ingrese los "
            "datos de la tarjeta — transparencia de descuentos (Épica 4, Regla 7)."
        ),
        "screenshot": SCREENSHOTS / "epic4-reservas-Epica-4---R-15240-e-frecuente-automaticamente-chromium" / "test-finished-1.png",
    },
    {
        "epic": "Épica 4 — Reservas",
        "num": 3,
        "title": "Validar restricciones de cupos disponibles",
        "given": "Paquete tiene solo 2 cupos disponibles",
        "when": "Usuario intenta reservar 5 pasajeros",
        "then": 'El sistema rechaza la reserva y muestra error "cupos insuficientes"',
        "description": (
            "Prueba el límite de capacidad (BookingService, Regla 4): un paquete nunca puede "
            "aceptar más pasajeros que cupos disponibles. La aplicación real implementa este "
            "rechazo en dos capas: el formulario de reserva bloquea el envío y muestra el mensaje "
            "de inmediato (prevención de errores, Heurística 5 de Nielsen — ver Documento 3), y el "
            "backend revalida la misma regla de forma independiente si la petición llegara a "
            "enviarse igual (defensa en profundidad)."
        ),
        "screenshot": SCREENSHOTS / "epic4-reservas-Epica-4---R-681ed-ciones-de-cupos-disponibles-chromium" / "test-finished-1.png",
    },
    {
        "epic": "Épica 5 — Pagos",
        "num": 1,
        "title": "Procesar pago exitoso y confirmar reserva",
        "given": "Usuario con reserva PENDING válida (monto = total_reserva)",
        "when": "Ingresa datos de tarjeta válidos y confirma pago",
        "then": "Sistema registra pago APPROVED, cambia reserva a CONFIRMED, muestra comprobante",
        "description": (
            "Cubre el cierre exitoso del flujo de compra (Épica 5, Reglas 6 y 10): con datos de "
            "tarjeta simulada válidos y un monto que coincide exactamente con el total de la "
            "reserva, el pago se registra como APPROVED, la reserva pasa de PENDING a CONFIRMED "
            "en la misma transacción, y el cliente ve un comprobante con los datos del viaje, el "
            "monto pagado y los últimos 4 dígitos de la tarjeta."
        ),
        "screenshot": SCREENSHOTS / "epic5-pagos-Epica-5---Pago-c9e7c-exitoso-y-confirmar-reserva-chromium" / "test-finished-1.png",
    },
    {
        "epic": "Épica 5 — Pagos",
        "num": 2,
        "title": "Rechazar tarjeta expirada",
        "given": "Usuario intenta pagar con tarjeta expirada (MM/YY < hoy)",
        "when": "Ingresa fecha de expiración pasada y confirma pago",
        "then": 'Sistema rechaza pago y muestra error "tarjeta expirada"',
        "description": (
            "Valida una de las reglas de formato de tarjeta (Épica 5, Regla 9) que no puede "
            "expresarse con una expresión regular: que el mes/año de expiración no sea una fecha "
            "ya pasada. Esta comprobación ocurre en el servicio (PaymentService), comparando "
            "contra la fecha actual del sistema, y debe rechazar el pago sin confirmar la reserva "
            "ni registrar ningún cobro."
        ),
        "screenshot": SCREENSHOTS / "epic5-pagos-Epica-5---Pago-aea7d-2-rechazar-tarjeta-expirada-chromium" / "test-finished-1.png",
    },
    {
        "epic": "Épica 5 — Pagos",
        "num": 3,
        "title": "Validar formato CVV de 3 dígitos",
        "given": "Usuario está en formulario de pago",
        "when": "Intenta ingresar CVV con formato inválido (menos de 3 dígitos)",
        "then": 'Sistema rechaza entrada y muestra error "CVV debe ser 3 dígitos"',
        "description": (
            "Valida el formato del CVV (Épica 5, Regla 9: @Pattern \\\\d{3}). En la aplicación "
            "real, el propio campo del formulario ya filtra cualquier carácter no numérico "
            "mientras el usuario escribe (prevención de errores en el cliente), así que el caso "
            "que efectivamente llega a validarse contra el backend es un CVV incompleto (menos de "
            "3 dígitos, pero numérico) — el sistema debe rechazarlo con un mensaje específico, no "
            "genérico."
        ),
        "screenshot": SCREENSHOTS / "epic5-pagos-Epica-5---Pago-15ac7-formato-de-CVV-de-3-digitos-chromium" / "test-finished-1.png",
    },
]


def set_cell_background(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_gherkin_table(doc, c):
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    labels = [("GIVEN", c["given"]), ("WHEN", c["when"]), ("THEN", c["then"])]
    for i, (label, text) in enumerate(labels):
        cell_label = table.cell(i, 0)
        cell_label.text = label
        cell_label.paragraphs[0].runs[0].bold = True
        cell_label.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell_label, "1E40AF")
        cell_label.width = Inches(1.1)
        table.cell(i, 1).width = Inches(5.5)
    return table


def heading(doc, text, level=1, color=PRIMARY):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = color
    return h


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Portada ---
    doc.add_paragraph().add_run().add_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Criterios de Aceptación")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Evaluación 3 — Épica 4 (Reservas) y Épica 5 (Pagos)")
    r.font.size = Pt(15)
    r.font.color.rgb = GRAY

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle2.add_run("App01Mingesoft — TravelAgency")
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Fecha: {date.today():%d/%m/%Y}\nFormato: Gherkin (Given / When / Then)")
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    doc.add_page_break()

    # --- Introducción ---
    heading(doc, "Introducción")
    doc.add_paragraph(
        "Este documento define los criterios de aceptación de las Épicas 4 (Reservas) y 5 "
        "(Pagos) de App01Mingesoft, en lenguaje Gherkin (Given/When/Then), como base para su "
        "posterior automatización con Playwright (ver Documento 2: Resultados de Pruebas "
        "Funcionales). Se definieron 3 criterios por épica, cubriendo en cada caso el flujo "
        "principal exitoso y al menos una regla de validación/rechazo relevante del negocio."
    )
    doc.add_paragraph(
        "Cada criterio incluye una descripción detallada de la regla de negocio que valida, "
        "referenciando el componente del backend responsable, y una captura de pantalla real de "
        "la aplicación en el estado que ese escenario ejercita."
    )

    doc.add_page_break()

    current_epic = None
    for c in CRITERIA:
        if c["epic"] != current_epic:
            current_epic = c["epic"]
            heading(doc, current_epic, level=1)

        heading(doc, f"Criterio {c['num']}: {c['title']}", level=2, color=RGBColor(0x1E, 0x3A, 0x8A))
        add_gherkin_table(doc, c)

        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("Descripción detallada: ")
        r.bold = True
        p.add_run(c["description"])

        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("Captura de referencia:")
        r.bold = True
        if c["screenshot"].exists():
            doc.add_picture(str(c["screenshot"]), width=Inches(5.8))
        else:
            doc.add_paragraph(f"[Captura no encontrada: {c['screenshot']}]")

        doc.add_paragraph()

    out_path = Path(__file__).parent / "Criterios_Aceptacion_Epic4_Epic5.docx"
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    build()
