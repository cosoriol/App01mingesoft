#!/usr/bin/env python3
"""Genera Documento 2 (Resultados_Pruebas_Funcionales.docx) para la Evaluacion 3.

Todos los datos (Gherkin, codigo, capturas, tiempos de ejecucion, resultado)
vienen de la corrida real de la suite Playwright en App01Mingesoft/tests/
(ver esa corrida en la conversacion / tests/report/index.html). No hay
numeros ni capturas inventados.
"""
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TESTS_DIR = Path(__file__).parent.parent.parent / "tests"
SCREENSHOTS = TESTS_DIR / "test-results"

PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
SUCCESS = RGBColor(0x15, 0x80, 0x3D)
GRAY = RGBColor(0x52, 0x52, 0x5B)

CRITERIA = [
    {
        "epic": "Épica 4 — Reservas",
        "num": 1,
        "title": "Crear reserva exitosa con validaciones",
        "given": "Usuario autenticado en la plataforma",
        "when": "Selecciona un paquete disponible y completa la reserva con datos válidos",
        "then": "El sistema crea la reserva, decrementa los cupos disponibles del paquete y avanza a la pantalla de pago",
        "spec_file": "specs/epic4-reservas.spec.js",
        "code": '''test('Criterio 1: crear reserva exitosa con validaciones', async ({ page }) => {
  // GIVEN: usuario autenticado en la plataforma
  const login = new LoginPage(page);
  await login.login(users.client1.email, helpers.TEST_PASSWORD);

  const detail = new PackageDetailPage(page);
  await detail.goto(pkgStandard.id);
  const slotsBefore = await detail.getAvailableSlotsText();

  // WHEN: selecciona un paquete disponible y completa la reserva con datos validos
  await detail.clickReserve();
  await page.waitForURL(new RegExp(`/booking/${pkgStandard.id}`));
  const booking = new BookingPage(page);
  await booking.setPassengerCount(2);
  await booking.continueToPayment();

  // THEN: el sistema crea la reserva (avanza a la pantalla de pago) ...
  await page.waitForURL(/\\/payment\\//);
  await expect(page.locator('h1')).toContainText('Confirmar Pago');

  // ... y decrementa los cupos disponibles del paquete
  await detail.goto(pkgStandard.id);
  const slotsAfter = await detail.getAvailableSlotsText();
  expect(slotsAfter).toBe(slotsBefore - 2);
});''',
        "screenshot": SCREENSHOTS / "epic4-reservas-Epica-4---R-16406-va-exitosa-con-validaciones-chromium" / "test-finished-1.png",
        "result": "PASS",
        "time": "3.9s",
        "note": None,
    },
    {
        "epic": "Épica 4 — Reservas",
        "num": 2,
        "title": "Aplicar descuento automáticamente en reserva (cliente frecuente)",
        "given": "Usuario es cliente frecuente (3+ reservas confirmadas)",
        "when": "Realiza una nueva reserva",
        "then": "El sistema aplica automáticamente el descuento por cliente frecuente (10%), visible en el resumen de pago",
        "spec_file": "specs/epic4-reservas.spec.js",
        "code": '''test('Criterio 2: aplicar descuento de cliente frecuente automaticamente', async ({ page }) => {
  // GIVEN: el usuario es cliente frecuente (>= 3 reservas CONFIRMED, armado en beforeAll)
  const login = new LoginPage(page);
  await login.login(users.frequentClient.email, helpers.TEST_PASSWORD);

  // WHEN: realiza una nueva reserva
  const detail = new PackageDetailPage(page);
  await detail.goto(pkgStandard.id);
  await detail.clickReserve();
  await page.waitForURL(new RegExp(`/booking/${pkgStandard.id}`));
  const booking = new BookingPage(page);
  await booking.setPassengerCount(1);
  await booking.continueToPayment();
  await page.waitForURL(/\\/payment\\//);

  // THEN: el sistema aplica automaticamente el descuento por cliente frecuente (10%),
  // visible en el resumen de pago antes de ingresar la tarjeta
  await expect(page.locator('body')).toContainText('Cliente frecuente');
  await expect(page.locator('body')).toContainText('10%');
});''',
        "screenshot": SCREENSHOTS / "epic4-reservas-Epica-4---R-15240-e-frecuente-automaticamente-chromium" / "test-finished-1.png",
        "result": "PASS",
        "time": "2.2s",
        "note": (
            "Precondición armada por API en beforeAll (helpers.ensureFrequentClient): se crean y "
            "pagan 3 reservas previas del usuario de prueba ANTES de correr el escenario por UI, "
            "para no depender de que ya existan datos de una corrida anterior."
        ),
    },
    {
        "epic": "Épica 4 — Reservas",
        "num": 3,
        "title": "Validar restricciones de cupos disponibles",
        "given": "Paquete tiene solo 2 cupos disponibles",
        "when": "Usuario intenta reservar 5 pasajeros",
        "then": "El sistema rechaza la reserva y muestra un mensaje indicando el rango válido de pasajeros",
        "spec_file": "specs/epic4-reservas.spec.js",
        "code": '''test('Criterio 3: validar restricciones de cupos disponibles', async ({ page }) => {
  // GIVEN: el paquete tiene solo 2 cupos disponibles
  const login = new LoginPage(page);
  await login.login(users.client3.email, helpers.TEST_PASSWORD);

  const detail = new PackageDetailPage(page);
  await detail.goto(pkgLowSlots.id);
  const slotsBefore = await detail.getAvailableSlotsText();
  expect(slotsBefore).toBe(2);

  // WHEN: el usuario intenta reservar 5 pasajeros
  await detail.clickReserve();
  await page.waitForURL(new RegExp(`/booking/${pkgLowSlots.id}`));
  const booking = new BookingPage(page);
  await booking.setPassengerCount(5);

  // THEN: el sistema rechaza la reserva. La UI la bloquea en el propio
  // formulario (Heuristica 5 de Nielsen: prevencion de errores, ver
  // BookingPage.js "validCount") en vez de esperar el rechazo del backend:
  // el boton queda deshabilitado y se muestra el mensaje de cupos.
  await expect(booking.errorMessage).toContainText(`entre 1 y ${slotsBefore}`);
  await expect(booking.continueButton).toBeDisabled();
});''',
        "screenshot": SCREENSHOTS / "epic4-reservas-Epica-4---R-681ed-ciones-de-cupos-disponibles-chromium" / "test-finished-1.png",
        "result": "PASS",
        "time": "2.0s",
        "note": (
            "Nota de diseño: el rechazo ocurre en el propio formulario (el botón “Continuar al "
            "pago” queda deshabilitado y aparece el mensaje de validación) en vez de esperar un "
            "error del servidor — es una aplicación correcta de la Heurística 5 de Nielsen "
            "(prevención de errores): el sistema impide la acción inválida antes de que ocurra, en "
            "lugar de dejarla ocurrir y luego mostrar un error. El backend (BookingService, REGLA 4) "
            "aplica la misma regla como segunda capa de defensa si se saltara la validación del "
            "cliente."
        ),
    },
    {
        "epic": "Épica 5 — Pagos",
        "num": 1,
        "title": "Procesar pago exitoso y confirmar reserva",
        "given": "Usuario con reserva PENDING válida (monto = total de la reserva)",
        "when": "Ingresa datos de tarjeta válidos y confirma el pago",
        "then": "El sistema registra el pago como APPROVED, cambia la reserva a CONFIRMED y muestra el comprobante",
        "spec_file": "specs/epic5-pagos.spec.js",
        "code": '''test('Criterio 1: procesar pago exitoso y confirmar reserva', async ({ page }) => {
  const user = await helpers.ensureUser(users.paymentClient1.fullName, users.paymentClient1.email);

  // GIVEN: usuario con reserva PENDING valida (monto = total de la reserva)
  const booking = await helpers.createBookingViaApi(user.id, pkgStandard.id, 1);

  const login = new LoginPage(page);
  await login.login(users.paymentClient1.email, helpers.TEST_PASSWORD);
  await page.goto(`/payment/${booking.id}`);

  // WHEN: ingresa datos de tarjeta validos y confirma el pago
  const paymentPage = new PaymentPage(page);
  await paymentPage.fillCard({
    holderName: 'Playwright Pago Uno',
    cardNumber: '4111111111111111',
    expirationDate: '1229', // se auto-formatea a 12/29
    cvv: '123',
  });
  await paymentPage.pay();

  // THEN: el sistema registra el pago (APPROVED), la reserva pasa a CONFIRMED,
  // y se muestra el comprobante
  await page.waitForURL(new RegExp(`/payment-confirmation/${booking.id}`));
  await expect(page.locator('h1')).toContainText('¡Pago realizado exitosamente!');
  await expect(page.locator('body')).toContainText('Confirmada');
});''',
        "screenshot": SCREENSHOTS / "epic5-pagos-Epica-5---Pago-c9e7c-exitoso-y-confirmar-reserva-chromium" / "test-finished-1.png",
        "result": "PASS",
        "time": "1.7s",
        "note": None,
    },
    {
        "epic": "Épica 5 — Pagos",
        "num": 2,
        "title": "Rechazar tarjeta expirada",
        "given": "Usuario intenta pagar con una tarjeta cuya fecha de expiración ya pasó (MM/YY < hoy)",
        "when": "Ingresa la fecha de expiración vencida y confirma el pago",
        "then": 'El sistema rechaza el pago y muestra el error "Card expiration date has already passed"',
        "spec_file": "specs/epic5-pagos.spec.js",
        "code": '''test('Criterio 2: rechazar tarjeta expirada', async ({ page }) => {
  const user = await helpers.ensureUser(users.paymentClient2.fullName, users.paymentClient2.email);
  const booking = await helpers.createBookingViaApi(user.id, pkgStandard.id, 1);

  const login = new LoginPage(page);
  await login.login(users.paymentClient2.email, helpers.TEST_PASSWORD);
  await page.goto(`/payment/${booking.id}`);

  // GIVEN/WHEN: intenta pagar con una tarjeta cuya fecha de expiracion ya paso
  const paymentPage = new PaymentPage(page);
  await paymentPage.fillCard({
    holderName: 'Playwright Pago Dos',
    cardNumber: '4111111111111111',
    expirationDate: '0120', // -> 01/20, Enero 2020: vencida
    cvv: '123',
  });
  await paymentPage.pay();

  // THEN: el sistema rechaza el pago y muestra el error de tarjeta expirada
  await expect(paymentPage.errorMessage).toContainText('expiration date has already passed');
  // La reserva sigue PENDING: el pago rechazado nunca debe confirmarla.
  await expect(page).toHaveURL(new RegExp(`/payment/${booking.id}`));
});''',
        "screenshot": SCREENSHOTS / "epic5-pagos-Epica-5---Pago-aea7d-2-rechazar-tarjeta-expirada-chromium" / "test-finished-1.png",
        "result": "PASS",
        "time": "1.5s",
        "note": None,
    },
    {
        "epic": "Épica 5 — Pagos",
        "num": 3,
        "title": "Validar formato de CVV de 3 dígitos",
        "given": "Usuario está en el formulario de pago",
        "when": "Ingresa un CVV incompleto (menos de 3 dígitos) y confirma el pago",
        "then": 'El sistema rechaza el pago y muestra el error "CVV must be exactly 3 digits"',
        "spec_file": "specs/epic5-pagos.spec.js",
        "code": '''test('Criterio 3: validar formato de CVV de 3 digitos', async ({ page }) => {
  const user = await helpers.ensureUser(users.paymentClient3.fullName, users.paymentClient3.email);
  const booking = await helpers.createBookingViaApi(user.id, pkgStandard.id, 1);

  const login = new LoginPage(page);
  await login.login(users.paymentClient3.email, helpers.TEST_PASSWORD);
  await page.goto(`/payment/${booking.id}`);

  const paymentPage = new PaymentPage(page);

  // GIVEN/WHEN: intenta ingresar un CVV con formato invalido.
  // El propio campo filtra letras en cada tecla (solo deja digitos, ver
  // PaymentPage.js) -- eso ya es prevencion de errores (Heuristica 5); el
  // caso que SI llega al backend es un CVV incompleto (menos de 3 digitos).
  await paymentPage.fillCard({
    holderName: 'Playwright Pago Tres',
    cardNumber: '4111111111111111',
    expirationDate: '1229',
    cvv: '12',
  });
  await paymentPage.pay();

  // THEN: el sistema rechaza el pago y muestra el error de CVV invalido
  await expect(paymentPage.errorMessage).toContainText('CVV must be exactly 3 digits');
});''',
        "screenshot": SCREENSHOTS / "epic5-pagos-Epica-5---Pago-15ac7-formato-de-CVV-de-3-digitos-chromium" / "test-finished-1.png",
        "result": "PASS",
        "time": "1.1s",
        "note": (
            "Hallazgo durante la automatización (bug real, no del test): las letras se filtran "
            "automáticamente al escribir en el campo CVV (PaymentPage.js), por lo que nunca llegan "
            "al backend — es prevención de errores (Heurística 5), no una validación probada aquí. "
            "El caso que sí ejercita la validación del servidor es un CVV incompleto (2 dígitos). "
            "Al automatizar este criterio se descubrió que el backend respondía con el mensaje "
            "genérico “Validation failed” en vez del mensaje específico exigido por este "
            "criterio: el frontend no leía el campo “details” de la respuesta de error. Se "
            "corrigió agregando un helper getErrorMessage() centralizado (services/api.js) y "
            "aplicándolo en las 11 páginas afectadas, para que el sistema cumpliera el criterio tal "
            "como está redactado, en vez de debilitar la aserción del test."
        ),
    },
]


def set_cell_background(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F7FB")
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rPr.append(rFonts)
    # Borde alrededor del bloque
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "CBD5E1")
        pBdr.append(el)
    pPr.append(pBdr)
    return p


def add_gherkin_table(doc, criterio):
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    labels = [("GIVEN", criterio["given"]), ("WHEN", criterio["when"]), ("THEN", criterio["then"])]
    for i, (label, text) in enumerate(labels):
        cell_label = table.cell(i, 0)
        cell_label.text = label
        cell_label.paragraphs[0].runs[0].bold = True
        cell_label.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell_label, "1E40AF")
        cell_label.width = Inches(1.1)

        cell_text = table.cell(i, 1)
        cell_text.text = text
        cell_text.width = Inches(5.5)
    return table


def add_result_badge(doc, criterio):
    p = doc.add_paragraph()
    run = p.add_run(f"  Resultado: {criterio['result']}  ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(11)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "15803D" if criterio["result"] == "PASS" else "B91C1C")
    pPr.append(shd)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Tiempo de ejecución: {criterio['time']}")
    r2.italic = True
    r2.font.color.rgb = GRAY
    r2.font.size = Pt(10)


def build():
    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Portada ---
    doc.add_paragraph().add_run().add_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Resultados de Pruebas Funcionales")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Evaluación 3 — Épica 4 (Reservas) y Épica 5 (Pagos)")
    r.font.size = Pt(15)
    r.font.color.rgb = GRAY

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle2.add_run("App01Mingesoft — TravelAgency")
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"Automatizado con Playwright (@playwright/test)\n"
        f"Fecha de ejecución: {date.today():%d/%m/%Y}\n"
        f"Suite: App01Mingesoft/tests/"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    doc.add_page_break()

    # --- Resumen ejecutivo ---
    h = doc.add_heading("Resumen ejecutivo", level=1)
    h.runs[0].font.color.rgb = PRIMARY

    doc.add_paragraph(
        "Se automatizaron los 6 criterios de aceptación definidos para las Épicas 4 (Reservas) y "
        "5 (Pagos), en lenguaje Gherkin, usando Playwright contra la interfaz real de "
        "App01Mingesoft (sin mocks). Los datos de arranque de cada escenario (usuarios, paquetes, "
        "historial de reservas previas) se preparan llamando directamente a la API real antes de "
        "cada prueba, para que cada corrida sea rápida y no dependa del estado dejado por una "
        "corrida anterior; el GIVEN/WHEN/THEN de cada criterio se ejecuta siempre contra la UI."
    )
    doc.add_paragraph(
        "Resultado: los 6 criterios pasan (6/6 PASS) contra el sistema real. Automatizar el "
        "Criterio 3 de Épica 5 (CVV inválido) encontró un bug real de UX — el sistema mostraba un "
        "mensaje de error genérico en vez del mensaje específico que exige el criterio — que se "
        "corrigió antes de dar por cumplida la evaluación, en vez de ajustar la prueba para que "
        "pasara igual. El detalle está en la sección de ese criterio."
    )

    summary_table = doc.add_table(rows=1, cols=5)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    headers = ["#", "Épica", "Criterio", "Resultado", "Tiempo"]
    for i, htext in enumerate(headers):
        hdr[i].text = htext
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr[i], "1E40AF")

    for c in CRITERIA:
        row = summary_table.add_row().cells
        row[0].text = str(c["num"])
        row[1].text = c["epic"].split("—")[0].strip()
        row[2].text = c["title"]
        row[3].text = c["result"]
        row[3].paragraphs[0].runs[0].font.color.rgb = SUCCESS
        row[3].paragraphs[0].runs[0].bold = True
        row[4].text = c["time"]

    total_row = summary_table.add_row().cells
    total_row[0].merge(total_row[2])
    total_row[0].text = "TOTAL (6 criterios)"
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[3].text = "6/6 PASS"
    total_row[3].paragraphs[0].runs[0].bold = True
    total_row[3].paragraphs[0].runs[0].font.color.rgb = SUCCESS
    total_row[4].text = "14.5s"
    total_row[4].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # --- Detalle por criterio ---
    current_epic = None
    for c in CRITERIA:
        if c["epic"] != current_epic:
            current_epic = c["epic"]
            h = doc.add_heading(current_epic, level=1)
            h.runs[0].font.color.rgb = PRIMARY

        h2 = doc.add_heading(f"Criterio {c['num']}: {c['title']}", level=2)
        h2.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

        add_gherkin_table(doc, c)
        doc.add_paragraph()

        p = doc.add_paragraph()
        r = p.add_run(f"Código Playwright ({c['spec_file']}):")
        r.bold = True
        add_code_block(doc, c["code"])

        p = doc.add_paragraph()
        r = p.add_run("Captura de ejecución:")
        r.bold = True
        if c["screenshot"].exists():
            doc.add_picture(str(c["screenshot"]), width=Inches(6.0))
        else:
            doc.add_paragraph(f"[Captura no encontrada: {c['screenshot']}]")

        add_result_badge(doc, c)

        if c["note"]:
            p = doc.add_paragraph()
            r = p.add_run("Nota: ")
            r.bold = True
            r.font.size = Pt(9.5)
            r2 = p.add_run(c["note"])
            r2.italic = True
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = GRAY

        doc.add_paragraph()

    out_path = Path(__file__).parent / "Resultados_Pruebas_Funcionales.docx"
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    build()
