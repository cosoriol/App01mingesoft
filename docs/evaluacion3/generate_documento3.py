#!/usr/bin/env python3
"""Genera Documento 3 (Usabilidad_Heuristicas_Nielsen.docx) para la Evaluacion 3.

Evaluacion honesta de las 10 heuristicas de Nielsen contra la aplicacion REAL
(codigo fuente + capturas de pantalla reales, App01Mingesoft/frontend). No
todas las heuristicas se cumplen al 100%: donde no, se documenta como
"Parcial" o "No" con la evidencia y una recomendacion concreta, en vez de
maquillar el resultado.
"""
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHOTS = Path(__file__).parent / "_nielsen_shots"
TEST_SHOTS = Path(__file__).parent.parent.parent / "tests" / "test-results"

PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
SUCCESS = RGBColor(0x15, 0x80, 0x3D)
WARNING = RGBColor(0xA1, 0x62, 0x07)
DANGER = RGBColor(0xB9, 0x1C, 0x1C)
GRAY = RGBColor(0x52, 0x52, 0x5B)

COMPLIANCE_COLOR = {"Sí": SUCCESS, "Parcial": WARNING, "No": DANGER}
COMPLIANCE_FILL = {"Sí": "15803D", "Parcial": "A16207", "No": "B91C1C"}

HEURISTICS = [
    {
        "num": 1,
        "title": "Visibilidad del estado del sistema",
        "desc": (
            "El sistema debe mantener informado al usuario sobre lo que está pasando, con "
            "retroalimentación apropiada en un tiempo razonable."
        ),
        "impl": (
            "Cada página que carga datos async muestra un estado de carga explícito "
            '("Cargando paquetes...", "Cargando tus reservas...", "Cargando comprobante..."). '
            "Toda acción que toma tiempo deshabilita el botón y cambia su texto mientras corre "
            '("Ingresando...", "Creando reserva...", "Procesando pago...", "Guardando..."). El '
            "estado de cada reserva se muestra siempre como una insignia de color (Pendiente de "
            "pago / Confirmada / Cancelada / Expirada), visible tanto en “Mis Reservas” como en "
            "el panel de administración."
        ),
        "screenshot": SHOTS / "04-my-bookings.png",
        "compliance": "Sí",
    },
    {
        "num": 2,
        "title": "Coincidencia entre el sistema y el mundo real",
        "desc": (
            "El sistema debe hablar el lenguaje del usuario, con palabras y conceptos "
            "familiares, en vez de términos orientados al sistema."
        ),
        "impl": (
            "Toda la interfaz está en español, con vocabulario del dominio de viajes: "
            "“Paquete”, “Reserva”, “Cupos disponibles”, “Cliente frecuente”, “Comprobante”. Los "
            "estados técnicos del backend (PENDING/CONFIRMED/CANCELLED/EXPIRED) se traducen "
            "siempre a texto natural (“Pendiente de pago”, “Confirmada”...) antes de mostrarse; "
            "el usuario nunca ve un enum ni un código de error crudo."
        ),
        "screenshot": SHOTS / "02-home.png",
        "compliance": "Sí",
    },
    {
        "num": 3,
        "title": "Control del usuario y libertad",
        "desc": (
            "Los usuarios necesitan una “salida de emergencia” claramente marcada para "
            "abandonar un estado no deseado, sin tener que pasar por un proceso extendido."
        ),
        "impl": (
            "Existen acciones de cancelación explícitas y siempre visibles cuando aplican: "
            "“Cancelar reserva” (mientras está PENDING o CONFIRMED), “Cancelar edición” al editar "
            "un paquete, “Cancelar paquete” / “Reactivar” para un administrador. No existe, en "
            "cambio, una función de deshacer tras confirmar una acción, y — limitación más "
            "notable — no hay una página de autorregistro (`/register`) en el frontend: un "
            "cliente nuevo no puede crear su propia cuenta por sí mismo, solo un administrador "
            "puede darla de alta hoy. Esto reduce la libertad del usuario en el punto de entrada "
            "al sistema."
        ),
        "screenshot": SHOTS / "03-admin-packages.png",
        "compliance": "Parcial",
        "recommendation": (
            "Agregar una página de autorregistro (`/register` en el frontend, consumiendo el "
            "endpoint ya existente `POST /api/auth/register`) para que un cliente nuevo pueda "
            "crear su cuenta sin depender de un administrador."
        ),
    },
    {
        "num": 4,
        "title": "Consistencia y estándares",
        "desc": (
            "Los usuarios no deberían tener que preguntarse si diferentes palabras, situaciones "
            "o acciones significan lo mismo. Seguir las convenciones de la plataforma."
        ),
        "impl": (
            "Los mismos componentes visuales (botones `btn-primary`/`btn-danger`/`btn-secondary`, "
            "tarjetas `summary-card`, campos `filter-group`, insignias de estado con la misma "
            "paleta de colores) se repiten sin variación en Home, Reservas, Pagos, Reportes y los "
            "paneles de administración. La barra de navegación y el pie de página son idénticos "
            "en todas las rutas."
        ),
        "screenshot": SHOTS / "03-admin-packages.png",
        "compliance": "Sí",
    },
    {
        "num": 5,
        "title": "Prevención de errores",
        "desc": (
            "Mejor que un buen mensaje de error es un diseño cuidadoso que prevenga que el "
            "problema ocurra en primer lugar."
        ),
        "impl": (
            "Evidencia concreta encontrada y verificada durante la automatización de la Épica 4 "
            "(ver Documento 2, Criterio 3): el formulario de reserva deshabilita el botón "
            "“Continuar al pago” y muestra el rango válido de pasajeros ANTES de intentar enviar "
            "la reserva, en vez de esperar el rechazo del servidor. En el formulario de pago, el "
            "campo de tarjeta y la fecha de expiración se autoformatean mientras se escribe, y el "
            "campo CVV filtra cualquier carácter no numérico tecla por tecla — nunca deja escribir "
            "una letra."
        ),
        "screenshot": TEST_SHOTS / "epic4-reservas-Epica-4---R-681ed-ciones-de-cupos-disponibles-chromium" / "test-finished-1.png",
        "compliance": "Sí",
    },
    {
        "num": 6,
        "title": "Reconocimiento en lugar de recuerdo",
        "desc": (
            "Minimizar la carga de memoria del usuario haciendo visibles objetos, acciones y "
            "opciones. El usuario no debería tener que recordar información de una parte de la "
            "interfaz a otra."
        ),
        "impl": (
            "Cada tarjeta de paquete muestra toda la información relevante para decidir "
            "(destino, precio, fechas, cupos) sin requerir navegar a otra pantalla. Los campos de "
            "búsqueda incluyen ejemplos como placeholder (“Ej: Machu Picchu”). El resumen de pago "
            "repite el detalle completo de la reserva (paquete, pasajeros, descuentos, total) "
            "junto al formulario de tarjeta, para que el usuario no tenga que recordar lo que "
            "reservó en la pantalla anterior."
        ),
        "screenshot": SHOTS / "02-home.png",
        "compliance": "Sí",
    },
    {
        "num": 7,
        "title": "Flexibilidad y eficiencia de uso",
        "desc": (
            "Los aceleradores — invisibles para el usuario novato — pueden acelerar la "
            "interacción para el usuario experto, de forma que el sistema atienda tanto a "
            "usuarios inexpertos como experimentados."
        ),
        "impl": (
            "No se encontraron atajos de teclado, modo de búsqueda avanzada, ni forma de "
            "personalizar/acelerar flujos repetitivos en el código del frontend. Los filtros de "
            "búsqueda visibles en la página principal cubren solo destino y rango de precio; el "
            "backend sí soporta filtrar además por fecha, tipo de viaje y temporada "
            "(`GET /api/packages/search`), pero esos filtros no están expuestos en la UI."
        ),
        "screenshot": SHOTS / "02-home.png",
        "compliance": "No",
        "recommendation": (
            "Exponer en el formulario de búsqueda los filtros de fecha/tipo de viaje/temporada "
            "que el backend ya acepta (no requiere cambios de API, solo de UI)."
        ),
    },
    {
        "num": 8,
        "title": "Diseño estético y minimalista",
        "desc": (
            "Las interfaces no deberían contener información irrelevante o rara vez necesaria. "
            "Cada unidad extra de información compite con las unidades relevantes."
        ),
        "impl": (
            "El diseño usa una paleta de colores acotada (azul primario, verde de éxito, rojo de "
            "peligro/cancelación, grises neutros), tipografía consistente, y tarjetas con "
            "espaciado uniforme. No hay elementos decorativos sin función; cada dato mostrado "
            "(precio, fechas, cupos, estado) es información que el usuario necesita para decidir "
            "o hacer seguimiento."
        ),
        "screenshot": SHOTS / "02-home.png",
        "compliance": "Sí",
    },
    {
        "num": 9,
        "title": "Ayuda y documentación",
        "desc": (
            "Aunque es mejor que el sistema pueda usarse sin documentación, puede ser necesario "
            "proveer ayuda. Esta información debe ser fácil de buscar, centrada en la tarea del "
            "usuario, y no demasiado extensa."
        ),
        "impl": (
            "No se encontraron tooltips, textos de ayuda contextual, ni una sección de FAQ/ayuda "
            "en ninguna página del frontend. Los únicos textos orientativos son los placeholders "
            "de los formularios y los mensajes de validación puntuales."
        ),
        "screenshot": SHOTS / "02-home.png",
        "compliance": "No",
        "recommendation": (
            "Agregar tooltips breves en los campos menos evidentes (ej. qué cuenta como "
            "“cliente frecuente”, cómo se calculan los descuentos combinados) y una página "
            "simple de preguntas frecuentes."
        ),
    },
    {
        "num": 10,
        "title": "Ayuda a reconocer, diagnosticar y recuperarse de errores",
        "desc": (
            "Los mensajes de error deben expresarse en lenguaje claro (sin códigos), indicar "
            "precisamente el problema, y sugerir una solución de forma constructiva."
        ),
        "impl": (
            "Los mensajes de error son específicos y accionables: “La cantidad de pasajeros debe "
            "ser entre 1 y 2”, “CVV must be exactly 3 digits”, “Card expiration date has already "
            "passed”. Durante esta evaluación se corrigió un bug real que afectaba esta "
            "heurística: el backend distingue errores de regla de negocio de errores de "
            "validación de formato, y estos últimos quedaban ocultos detrás de un mensaje "
            "genérico “Validation failed” en 11 páginas del frontend — se agregó un helper "
            "`getErrorMessage()` centralizado para que el mensaje específico siempre llegue al "
            "usuario (ver Documento 2, Criterio 3 de Épica 5)."
        ),
        "screenshot": TEST_SHOTS / "epic5-pagos-Epica-5---Pago-15ac7-formato-de-CVV-de-3-digitos-chromium" / "test-finished-1.png",
        "compliance": "Sí",
    },
]


def set_cell_background(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def heading(doc, text, level=1, color=PRIMARY):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = color
    return h


def add_compliance_badge(doc, compliance):
    p = doc.add_paragraph()
    run = p.add_run(f"  Cumplimiento: {compliance}  ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(11)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), COMPLIANCE_FILL[compliance])
    pPr.append(shd)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Portada ---
    doc.add_paragraph().add_run().add_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Usabilidad — Heurísticas de Nielsen")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Evaluación 3 — Pruebas No Funcionales: Usabilidad")
    r.font.size = Pt(15)
    r.font.color.rgb = GRAY

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle2.add_run("App01Mingesoft — TravelAgency")
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Fecha de evaluación: {date.today():%d/%m/%Y}")
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    doc.add_page_break()

    # --- Resumen ---
    heading(doc, "Resumen")
    doc.add_paragraph(
        "Se evaluaron las 10 heurísticas de usabilidad de Jakob Nielsen contra la aplicación "
        "real (código fuente del frontend + navegación real capturada con Playwright), no contra "
        "una descripción idealizada. El resultado se documenta con honestidad: donde la "
        "aplicación no cumple completamente una heurística, se marca como “Parcial” o “No”, con "
        "la evidencia concreta y una recomendación puntual — en vez de calificar todo como "
        "cumplido."
    )

    summary = doc.add_table(rows=1, cols=3)
    summary.style = "Table Grid"
    hdr = summary.rows[0].cells
    for i, h in enumerate(["#", "Heurística", "Cumplimiento"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr[i], "1E40AF")
    for h in HEURISTICS:
        row = summary.add_row().cells
        row[0].text = str(h["num"])
        row[1].text = h["title"]
        row[2].text = h["compliance"]
        row[2].paragraphs[0].runs[0].bold = True
        row[2].paragraphs[0].runs[0].font.color.rgb = COMPLIANCE_COLOR[h["compliance"]]

    si_count = sum(1 for h in HEURISTICS if h["compliance"] == "Sí")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"Resultado global: {si_count}/10 heurísticas cumplidas completamente, "
                  f"{sum(1 for h in HEURISTICS if h['compliance']=='Parcial')} parciales, "
                  f"{sum(1 for h in HEURISTICS if h['compliance']=='No')} no cumplidas.")
    r.bold = True

    doc.add_page_break()

    for h in HEURISTICS:
        heading(doc, f"Heurística {h['num']}: {h['title']}", level=1)

        p = doc.add_paragraph()
        r = p.add_run("Descripción: ")
        r.bold = True
        r2 = p.add_run(h["desc"])
        r2.italic = True

        p = doc.add_paragraph()
        r = p.add_run("Implementación en App01Mingesoft: ")
        r.bold = True
        p.add_run(h["impl"])

        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("Evidencia (captura real):")
        r.bold = True
        if h["screenshot"].exists():
            doc.add_picture(str(h["screenshot"]), width=Inches(5.6))
        else:
            doc.add_paragraph(f"[Captura no encontrada: {h['screenshot']}]")

        add_compliance_badge(doc, h["compliance"])

        if "recommendation" in h:
            p = doc.add_paragraph()
            r = p.add_run("Recomendación: ")
            r.bold = True
            r.font.size = Pt(9.5)
            r2 = p.add_run(h["recommendation"])
            r2.italic = True
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = GRAY

        doc.add_paragraph()

    out_path = Path(__file__).parent / "Usabilidad_Heuristicas_Nielsen.docx"
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    build()
