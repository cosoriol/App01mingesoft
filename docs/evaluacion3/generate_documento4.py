#!/usr/bin/env python3
"""Genera Documento 4 (Resultado_Cuestionario_SUS.docx) para la Evaluacion 3.

A diferencia de los otros 5 documentos, este NO puede generarse completo de
forma automatica: el cuestionario SUS exige respuestas de al menos 5
personas reales, y no existe una forma honesta de producir eso sin datos
humanos de verdad. Este script arma la PLANTILLA (preguntas, formula de
calculo, tabla vacia) lista para completar en cuanto existan respuestas
reales -- ver sus-survey/index.html para la encuesta real que se comparte
con los participantes.
"""
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
WARNING = RGBColor(0xA1, 0x62, 0x07)
GRAY = RGBColor(0x52, 0x52, 0x5B)

QUESTIONS = [
    "Creo que me gustaría usar este sistema frecuentemente.",
    "El sistema es innecesariamente complejo.",
    "El sistema es fácil de usar.",
    "Creo que necesitaría ayuda técnica para usar este sistema.",
    "Las funciones del sistema están bien integradas entre sí.",
    "Siento que hay demasiada inconsistencia en este sistema.",
    "Creo que la mayoría de la gente aprendería a usar este sistema muy rápidamente.",
    "El sistema es muy incómodo de usar.",
    "Me siento muy confiado usando el sistema.",
    "Necesité aprender muchas cosas antes de poder empezar a usar este sistema.",
]

MIN_PARTICIPANTS = 5


def set_cell_background(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def heading(doc, text, level=1, color=PRIMARY):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = color
    return h


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Portada ---
    doc.add_paragraph().add_run().add_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Resultado del Cuestionario SUS")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Evaluación 3 — Pruebas No Funcionales: Usabilidad (System Usability Scale)")
    r.font.size = Pt(14)
    r.font.color.rgb = GRAY

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle2.add_run("App01Mingesoft — TravelAgency")
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Plantilla generada: {date.today():%d/%m/%Y}")
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    doc.add_page_break()

    # --- Aviso de estado ---
    warn_p = doc.add_paragraph()
    warn_pPr = warn_p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FBF1E2")
    warn_pPr.append(shd)
    r = warn_p.add_run(
        "⚠ ESTADO: PLANTILLA — pendiente de respuestas reales.\n"
    )
    r.bold = True
    r.font.color.rgb = WARNING
    r2 = warn_p.add_run(
        f"El cuestionario SUS exige respuestas de al menos {MIN_PARTICIPANTS} personas reales; no "
        "es un dato que pueda generarse automáticamente sin falsear la evaluación. Este documento "
        "trae la encuesta, la fórmula de cálculo ya explicada, y la tabla de resultados vacía, "
        "lista para completar en cuanto se recojan las respuestas reales. La sección "
        "“Cómo completar este documento”, más abajo, explica el proceso."
    )
    r2.font.color.rgb = WARNING

    doc.add_page_break()

    # --- Como completar ---
    heading(doc, "Cómo completar este documento")
    doc.add_paragraph(
        "1. Comparte el enlace de la encuesta (ver más abajo) con al menos 5 personas: pueden ser "
        "docentes, compañeros, familiares o usuarios externos — personas que no hayan participado "
        "en el desarrollo de la aplicación, para que la evaluación sea independiente."
    )
    doc.add_paragraph(
        "2. Cada persona responde las 10 preguntas (toma menos de 3 minutos) y, al final, la "
        "encuesta le muestra su puntaje SUS individual y un botón para copiarlo o enviarlo por "
        "email."
    )
    doc.add_paragraph(
        "3. Reúne los resultados de las 5+ personas y complétalos en la tabla de la sección "
        "“Tabulación de resultados” de este documento (fila por participante, con su puntaje "
        "SUS)."
    )
    doc.add_paragraph(
        "4. Calcula el promedio de la columna “Score SUS” y complétalo en la fila TOTAL. Compáralo "
        "contra el umbral de 75 (ver sección “Interpretación”) para la conclusión final."
    )

    p = doc.add_paragraph()
    r = p.add_run("Encuesta (compartir este enlace): ")
    r.bold = True
    p.add_run(
        "el archivo fuente vive en App01Mingesoft/docs/evaluacion3/sus-survey/index.html — "
        "publicado como Artifact para compartir el enlace directamente con los participantes."
    )

    doc.add_page_break()

    # --- Las 10 preguntas ---
    heading(doc, "Las 10 preguntas del cuestionario")
    doc.add_paragraph(
        "Escala de 1 (muy en desacuerdo) a 5 (muy de acuerdo), respondida pensando en la "
        "experiencia recién vivida usando la aplicación."
    )
    for i, q in enumerate(QUESTIONS, start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(q)

    doc.add_page_break()

    # --- Calculo SUS ---
    heading(doc, "Cálculo del puntaje SUS")
    doc.add_paragraph(
        "El System Usability Scale (Brooke, 1986) combina las 10 respuestas en un único puntaje "
        "de 0 a 100 mediante esta fórmula estándar:"
    )
    for text in [
        "Preguntas impares (1, 3, 5, 7, 9): puntaje de la pregunta − 1",
        "Preguntas pares (2, 4, 6, 8, 10): 5 − puntaje de la pregunta",
        "Se suman los 10 valores ajustados (rango resultante: 0 a 40)",
        "El total se multiplica por 2,5, dando un puntaje final de 0 a 100",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
    doc.add_paragraph(
        "Este cálculo ya está implementado en la encuesta (index.html): cada participante ve su "
        "propio puntaje al terminar, no hace falta calcularlo a mano — solo transcribirlo a la "
        "tabla de abajo."
    )

    doc.add_page_break()

    # --- Tabulacion (vacia) ---
    heading(doc, "Tabulación de resultados")
    p = doc.add_paragraph()
    r = p.add_run(
        f"Completar con los {MIN_PARTICIPANTS}+ participantes reales. Filas de ejemplo no incluidas "
        "a propósito, para no confundir un placeholder con un dato real."
    )
    r.italic = True
    r.font.color.rgb = GRAY

    table = doc.add_table(rows=1, cols=13)
    table.style = "Table Grid"
    headers = ["Participante"] + [f"P{i}" for i in range(1, 11)] + ["Score SUS"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "1E40AF")

    for _ in range(MIN_PARTICIPANTS):
        row = table.add_row().cells
        for c in row:
            c.text = ""

    total_row = table.add_row().cells
    total_row[0].text = "PROMEDIO"
    total_row[0].paragraphs[0].runs[0].bold = True
    for c in total_row[1:12]:
        c.text = "—"
    total_row[12].text = ""

    doc.add_page_break()

    # --- Interpretacion ---
    heading(doc, "Interpretación")
    doc.add_paragraph(
        "Según la escala de Bangor, Kortum y Miller (2009), ampliamente usada para interpretar "
        "puntajes SUS:"
    )
    interp_table = doc.add_table(rows=1, cols=2)
    interp_table.style = "Table Grid"
    interp_table.rows[0].cells[0].text = "Puntaje SUS"
    interp_table.rows[0].cells[1].text = "Interpretación"
    for c in interp_table.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(c, "1E40AF")
    for score, label in [
        ("≥ 80.3", "Excelente"),
        ("68 – 80.2", "Bueno"),
        ("51 – 67.9", "Aceptable / regular"),
        ("< 51", "Deficiente — requiere revisión de usabilidad"),
    ]:
        row = interp_table.add_row().cells
        row[0].text = score
        row[1].text = label

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Umbral objetivo de esta evaluación: ")
    r.bold = True
    p.add_run("promedio ≥ 75 (recomendado por el profesor), correspondiente a la banda “Bueno”.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Conclusión: ")
    r.bold = True
    r.font.color.rgb = WARNING
    p.add_run("[Completar una vez calculado el promedio real de la tabla anterior]")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Mejoras sugeridas: ")
    r.bold = True
    p.add_run(
        "si el promedio queda por debajo del umbral, revisar primero las heurísticas de Nielsen "
        "marcadas como “Parcial” o “No” en el Documento 3 (autorregistro ausente, falta de "
        "filtros de búsqueda avanzados en la UI, falta de ayuda/FAQ) — son las brechas de "
        "usabilidad ya identificadas de forma independiente y probablemente correlacionan con "
        "puntajes SUS bajos en preguntas como P3 (facilidad de uso) o P7 (rapidez de aprendizaje)."
    )

    out_path = Path(__file__).parent / "Resultado_Cuestionario_SUS.docx"
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    build()
