#!/usr/bin/env python3
"""Genera Documento 6 (Resumen_Ejecutivo_Evaluacion3.docx) para la Evaluacion 3.

Sintesis de los 5 documentos anteriores. Reporta el estado real de cada
frente -- incluido lo que sigue pendiente (SUS necesita respuestas humanas
reales) -- en vez de marcar todo como completado.
"""
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
SUCCESS = RGBColor(0x15, 0x80, 0x3D)
WARNING = RGBColor(0xA1, 0x62, 0x07)
GRAY = RGBColor(0x52, 0x52, 0x5B)

STATUS_FILL = {"Completo": "15803D", "Pendiente": "A16207", "Parcial": "A16207"}


def set_cell_background(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def heading(doc, text, level=1, color=PRIMARY):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = color
    return h


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Portada ---
    doc.add_paragraph().add_run().add_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Resumen Ejecutivo")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Evaluación 3 — Pruebas Funcionales y No Funcionales")
    r.font.size = Pt(15)
    r.font.color.rgb = GRAY

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle2.add_run("App01Mingesoft — TravelAgency")
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Fecha: {date.today():%d/%m/%Y}")
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    doc.add_page_break()

    # --- Estado general ---
    heading(doc, "Estado general")
    doc.add_paragraph(
        "De los 6 documentos de la Evaluación 3, 5 están completos con datos reales verificados "
        "contra la aplicación en ejecución. El sexto (Documento 4, cuestionario SUS) tiene la "
        "encuesta lista y funcionando, pero su resultado depende de respuestas de personas reales "
        "que aún no se han recolectado — no es algo que pueda completarse sin ese paso humano."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Documento", "Contenido", "Estado"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr[i], "1E40AF")

    rows = [
        ("1. Criterios de Aceptación", "6 criterios Gherkin (Épica 4 y 5) con descripción y evidencia", "Completo"),
        ("2. Resultados Pruebas Funcionales", "Automatización Playwright: 6/6 criterios PASS", "Completo"),
        ("3. Usabilidad — Nielsen", "10 heurísticas evaluadas: 7 cumplidas, 2 parciales, 1 no cumplida", "Completo"),
        ("4. Resultado Cuestionario SUS", "Encuesta funcionando; tabla pendiente de 5+ respuestas reales", "Pendiente"),
        ("5. Pruebas de Rendimiento K6", "Load/Stress/Volume reales; cuello de botella identificado", "Completo"),
        ("6. Resumen Ejecutivo", "Este documento", "Completo"),
    ]
    for name, content, status in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = content
        row[2].text = status
        row[2].paragraphs[0].runs[0].bold = True
        row[2].paragraphs[0].runs[0].font.color.rgb = SUCCESS if status == "Completo" else WARNING

    doc.add_page_break()

    # --- Metricas clave ---
    heading(doc, "Métricas clave")

    heading(doc, "Pruebas funcionales", level=2, color=RGBColor(0x1E, 0x3A, 0x8A))
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("6/6 criterios de aceptación PASS (Épica 4: 3/3, Épica 5: 3/3), tiempo total de ejecución: 14.5s.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(
        "1 bug real encontrado y corregido durante la automatización (mensajes de validación "
        "genéricos en 11 páginas del frontend, en vez del mensaje específico exigido por el "
        "criterio) — el sistema se corrigió para cumplir el criterio tal como está redactado, en "
        "vez de debilitar la prueba."
    )

    heading(doc, "Usabilidad (heurísticas de Nielsen)", level=2, color=RGBColor(0x1E, 0x3A, 0x8A))
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("7/10 heurísticas cumplidas completamente, 2 parciales, 1 no cumplida.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(
        "Brechas identificadas: sin autorregistro de usuarios en el frontend (Heurística 3), sin "
        "filtros de búsqueda avanzados expuestos en la UI aunque el backend los soporta "
        "(Heurística 7), sin ayuda/FAQ contextual (Heurística 9)."
    )

    heading(doc, "Rendimiento (K6)", level=2, color=RGBColor(0x1E, 0x3A, 0x8A))
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Load test: 0% de errores reales hasta 200 usuarios concurrentes, P95 < 10 ms en todos los niveles.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Stress test: sin punto de quiebre hasta 150 usuarios concurrentes (P95 = 5.99 ms, 0% error).")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(
        "Volume test: degradación ~lineal con el volumen de datos — P95 de 8.26 ms (500 registros) "
        "a 110.42 ms (10.000 registros). Cuello de botella real identificado en "
        "ReportService.getSalesReport() (filtrado/agregación en memoria en vez de en SQL)."
    )

    doc.add_page_break()

    # --- Recomendaciones ---
    heading(doc, "Recomendaciones")

    recs = [
        (
            "Recolectar las respuestas SUS reales",
            "Compartir la encuesta (Documento 4 / sus-survey/index.html) con al menos 5 personas "
            "y completar la tabla de resultados — es el único punto abierto de toda la evaluación.",
        ),
        (
            "Optimizar ReportService.getSalesReport()",
            "Mover el filtro por estado y la agregación (SUM/COUNT) a una consulta @Query en vez "
            "de traer todas las reservas del rango a memoria. Es la causa raíz confirmada de la "
            "degradación de rendimiento con el volumen de datos (Documento 5).",
        ),
        (
            "Agregar autorregistro de usuarios",
            "Falta una página `/register` en el frontend que consuma el endpoint "
            "`POST /api/auth/register` ya existente — hoy solo un administrador puede crear "
            "cuentas nuevas (Documento 3, Heurística 3).",
        ),
        (
            "Exponer los filtros de búsqueda avanzados existentes",
            "El backend ya soporta filtrar por fecha, tipo de viaje y temporada "
            "(`GET /api/packages/search`); solo falta agregarlos al formulario de búsqueda del "
            "frontend (Documento 3, Heurística 7).",
        ),
        (
            "Agregar ayuda contextual mínima",
            "Tooltips en los campos menos evidentes (regla de cliente frecuente, cálculo de "
            "descuentos combinados) y una página simple de preguntas frecuentes (Documento 3, "
            "Heurística 9).",
        ),
    ]
    for title, detail in recs:
        p = doc.add_paragraph()
        r = p.add_run(f"{title}. ")
        r.bold = True
        p.add_run(detail)
        doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(
        "Ninguna de estas recomendaciones bloquea la entrega de la Evaluación 3: los 6 criterios "
        "funcionales pasan, el sistema resiste la carga y el estrés probados, y las brechas de "
        "usabilidad encontradas están acotadas y documentadas con evidencia concreta, no son "
        "fallas críticas del sistema."
    )
    r.italic = True

    out_path = Path(__file__).parent / "Resumen_Ejecutivo_Evaluacion3.docx"
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    build()
